# -*- coding: utf-8 -*-
"""
__title__ = ''
__author__ = 'Marc_lu'
__mtime__ = '2020-9-20'
# TODO: 遍历读取该目录下的文件， 将 .py 文件内容保存到word中
### how to use: 1. pip install python-docx
###             2. python main.py
"""

import os
from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_LINE_SPACING


def write2doc(filename):
    with open(filename, encoding='UTF-8') as f:
        for line in f.readlines():
            if line == '\n':  # 删除空行
                continue
            if line.lstrip().startswith('#'):  # 删除注释
                continue
            # if re.match(r'^\s+$', line):  # 使用正则表达式删除全是空格的空行
                # continue
            # 注释为''' '''的暂不考虑
            p.add_run(line)
            global codeNum
            codeNum += 1
            '''
            if codeNum == 3050:  # 保证打印出不大大超过与 60 页
                return
            '''


doc = Document()
p = doc.add_paragraph('')  # 增加一页
doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
doc.styles['Normal'].font.size = Pt(8)  # 设置字体的大小为 5 号字体
p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
paragraph_format = doc.styles['Normal'].paragraph_format
paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码
save_file = r'E:/Program/Underwater/FinalDesign/code.doc'
if os.path.exists(save_file):
    os.remove(save_file)
dir_path = r"E:/Program/Underwater/FinalDesign"
codeNum = 0
for root, dirs, files in os.walk(dir_path):
    for file in files:
        if file.endswith('.py'):
            code_file = os.path.join(root, file)
            write2doc(code_file)


doc.save(save_file)  # 不足60 页进行保存
print('all done')
print("代码行数：", codeNum)
